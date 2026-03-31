"""
TDD tests for document extraction (PDF, Word, .eml, .txt).
DoR: document_extractor module exists. DoD: extract_document_text supports .eml, .txt, .md, .pdf, .docx
"""
import pytest
from pathlib import Path
import tempfile
import sys

PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


@pytest.fixture
def sample_txt_path(tmp_path):
    """Create sample .txt file."""
    p = tmp_path / "sample.txt"
    p.write_text("Settlement offer. Case 26CV005596. Evidence: 40+ work orders. Deadline Feb 12.", encoding="utf-8")
    return p


@pytest.fixture
def sample_eml_path(tmp_path):
    """Create sample .eml file."""
    p = tmp_path / "sample.eml"
    p.write_text(
        "Subject: Settlement\nPursuant to N.C.G.S. § 42-42. Pro Se. Case No.: 26CV005596-590.",
        encoding="utf-8",
    )
    return p


class TestExtractTextFile:
    """TDD: Plain text files (.txt, .md, .eml, .html)"""

    def test_extract_txt_returns_content(self, sample_txt_path):
        from vibesthinker.document_extractor import extract_document_text
        content = extract_document_text(str(sample_txt_path))
        assert "Settlement" in content
        assert "26CV005596" in content

    def test_extract_eml_returns_content(self, sample_eml_path):
        from vibesthinker.document_extractor import extract_document_text
        content = extract_document_text(str(sample_eml_path))
        assert "N.C.G.S." in content
        assert "Pro Se" in content

    def test_extract_nonexistent_raises(self):
        from vibesthinker.document_extractor import extract_document_text
        with pytest.raises(FileNotFoundError):
            extract_document_text("/nonexistent/path.txt")


class TestExtractPdf:
    """TDD: PDF extraction (requires pypdf)"""

    def test_extract_pdf_returns_string(self):
        """Minimal PDF returns string (blank pages => placeholder)."""
        pytest.importorskip("pypdf")
        from vibesthinker.document_extractor import extract_document_text
        from pypdf import PdfWriter
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            pdf_path = f.name
        try:
            writer = PdfWriter()
            writer.add_blank_page(100, 100)
            with open(pdf_path, "wb") as f:
                writer.write(f)
            content = extract_document_text(pdf_path)
            assert isinstance(content, str)
            assert len(content) > 0
        finally:
            Path(pdf_path).unlink(missing_ok=True)



class TestExtractDocx:
    """TDD: Word .docx extraction (requires python-docx)"""

    def test_extract_docx_returns_content(self):
        docx = pytest.importorskip("docx")
        from docx import Document
        from vibesthinker.document_extractor import extract_document_text
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = f.name
        try:
            doc = Document()
            doc.add_paragraph("Settlement proposal. Case 26CV005596. Evidence attached.")
            doc.save(docx_path)
            content = extract_document_text(docx_path)
            assert "Settlement" in content
            assert "26CV005596" in content
        finally:
            Path(docx_path).unlink(missing_ok=True)

    def test_extract_docx_empty_doc_returns_string(self):
        """Empty .docx returns newline-separated string."""
        docx = pytest.importorskip("docx")
        from docx import Document
        from vibesthinker.document_extractor import extract_document_text
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = f.name
        try:
            doc = Document()
            doc.save(docx_path)
            content = extract_document_text(docx_path)
            assert isinstance(content, str)
        finally:
            Path(docx_path).unlink(missing_ok=True)


class TestExtractSupportedFormats:
    """TDD: supported_formats() returns expected extensions"""

    def test_supported_formats_includes_txt_eml_pdf_docx(self):
        from vibesthinker.document_extractor import supported_formats
        formats = supported_formats()
        assert ".txt" in formats
        assert ".eml" in formats
        assert ".md" in formats
        assert ".pdf" in formats
        assert ".docx" in formats
